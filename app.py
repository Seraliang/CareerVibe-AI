"""
CareerVibe — 深度分析版（Streamlit，非流式）
小 JSON 拉取雷达五维；三段正文各自 chat_completion 完成后一次性展示。
主区（三 Tab 下方）「面试预测」「面试复盘」为独立非流式任务，仅按钮触发。
"""

from __future__ import annotations 

import io
import json
import math
import os
import traceback
from pathlib import Path
from typing import Any

import pdfplumber
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI

_PROJECT_ROOT = Path(__file__).resolve().parent
load_dotenv(_PROJECT_ROOT / ".env")
load_dotenv()

RADAR_LABELS_CN = ["硬技能", "软实力", "行业匹配", "经验资历", "发展潜力"]
DEFAULT_RADAR_SCORES = [80, 80, 80, 80, 80]

SYSTEM_HEADHUNTER = """你是资深公关与传播行业猎头顾问，中文输出，专业、直接、可执行。
严格依据简历与 JD；证据不足时用一句话写清假设。"""

SYSTEM_INTERVIEW_MASTER = """你扮演 MBB（麦肯锡/贝恩/BCG）或 Google/Meta 等级别企业中的资深 HR Director / Hiring Bar Raiser：
面试语言为中文；风格冷静、结构化、高标准；善于基于候选人「真实履历细节」设计追问与示范回答。
输出必须可扫读：标题层级清晰，题与题之间用水平分割线（Markdown ---）分隔。
用户要求 **恰好 10 道** Masterclass 级题目：每题「满分范例回答」必须写深写透；**必须完整输出 Q1–Q10，不得中途省略或烂尾。**"""


def inject_deep_css() -> None:
    """Custom theme disabled — default Streamlit UI (sidebar debugging)."""
    pass


def extract_pdf_text(uploaded_file) -> str:
    raw = uploaded_file.getvalue()
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                parts.append(t.strip())
    return "\n\n".join(parts).strip()


def extract_docx_text(uploaded_file) -> str:
    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def extract_resume_text(uploaded_file) -> str:
    name = (getattr(uploaded_file, "name", None) or "").lower()
    if name.endswith(".docx"):
        return extract_docx_text(uploaded_file)
    return extract_pdf_text(uploaded_file)


def get_api_key() -> str:
    return (
        os.getenv("DEEPSEEK_API_KEY", "").strip()
        or os.getenv("OPENAI_API_KEY", "").strip()
    )


def get_base_url() -> str:
    raw = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com").strip()
    return raw or "https://api.deepseek.com"


def get_client() -> OpenAI:
    key = get_api_key()
    if not key:
        st.error(
            "未检测到 API 密钥。请在 `.env` 中设置 `DEEPSEEK_API_KEY`（或临时兼容 `OPENAI_API_KEY`）。"
        )
        st.stop()
    return OpenAI(api_key=key, base_url=get_base_url())


def get_model() -> str:
    return (
        os.getenv("DEEPSEEK_MODEL", "").strip()
        or os.getenv("OPENAI_MODEL", "").strip()
        or "deepseek-chat"
    )


def chat_completion(
    client: OpenAI,
    system: str,
    user: str,
    *,
    json_mode: bool = False,
    max_tokens: int = 2048,
) -> str:
    kwargs: dict[str, Any] = {
        "model": get_model(),
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.35,
        "max_tokens": max_tokens,
    }
    if json_mode:
        kwargs["response_format"] = {"type": "json_object"}
    resp = client.chat.completions.create(**kwargs)
    return (resp.choices[0].message.content or "").strip()


def clamp_int_score(v: object, fallback: int) -> int:
    try:
        x = int(float(v))
    except (TypeError, ValueError):
        return fallback
    if not math.isfinite(x):
        return fallback
    return max(0, min(100, x))


def normalize_radar_list(raw: object) -> list[int] | None:
    if not isinstance(raw, list) or len(raw) < 5:
        return None
    out: list[int] = []
    for i in range(5):
        out.append(clamp_int_score(raw[i], DEFAULT_RADAR_SCORES[i]))
    return out


def fetch_radar_scores_safe(
    client: OpenAI, resume: str, jd: str
) -> tuple[list[int], str | None]:
    """仅请求雷达 JSON；失败返回默认 80×5 与提示文案。"""
    user = f"""仅输出合法 JSON（不要 Markdown 围栏），且 **只能** 包含一个键 radar_scores：
{{ "radar_scores": [a,b,c,d,e] }}
五个整数 0-100，顺序固定为：{", ".join(RADAR_LABELS_CN)}。

--- 简历 ---
{resume[:8000]}

--- JD ---
{jd[:8000]}
"""
    try:
        raw = chat_completion(
            client, SYSTEM_HEADHUNTER, user, json_mode=True, max_tokens=400
        )
        data = json.loads(raw.strip())
        if not isinstance(data, dict):
            return list(DEFAULT_RADAR_SCORES), "雷达分数解析失败，已使用默认值 80。"
        arr = normalize_radar_list(data.get("radar_scores"))
        if arr is None:
            return list(DEFAULT_RADAR_SCORES), "雷达分数格式异常，已使用默认值 80。"
        return arr, None
    except Exception:  # noqa: BLE001
        return list(DEFAULT_RADAR_SCORES), "雷达分数接口异常，已使用默认值 80。"


def radar_list_to_dict(scores: list[int]) -> dict[str, int]:
    return {RADAR_LABELS_CN[i]: scores[i] for i in range(5)}


def overall_from_scores(scores: list[int]) -> int:
    if len(scores) != 5:
        return 80
    return int(round(sum(scores) / 5.0))


def build_radar_figure(scores: list[int]) -> go.Figure:
    d = radar_list_to_dict(scores)
    values = [d[k] for k in RADAR_LABELS_CN]
    theta = RADAR_LABELS_CN + [RADAR_LABELS_CN[0]]
    r = values + [values[0]]
    fig = go.Figure()
    fig.add_trace(
        go.Scatterpolar(
            r=r,
            theta=theta,
            fill="toself",
            fillcolor="rgba(15, 23, 42, 0.28)",
            line=dict(color="#0f172a", width=3),
            hovertemplate="%{theta}: %{r}<extra></extra>",
        )
    )
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 100],
                gridcolor="rgba(15,23,42,0.14)",
                linecolor="rgba(15,23,42,0.18)",
            ),
            angularaxis=dict(linecolor="rgba(15,23,42,0.22)"),
            bgcolor="rgba(0,0,0,0)",
        ),
        showlegend=False,
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(t=20, b=24, l=24, r=24),
        font=dict(color="#0f172a", size=11, family="sans-serif"),
    )
    return fig


def render_radar_top(scores: list[int], overall: int) -> None:
    try:
        c1, c2 = st.columns([0.88, 1.32])
        with c1:
            st.metric("综合匹配", f"{overall} / 100")
        with c2:
            st.plotly_chart(
                build_radar_figure(scores),
                use_container_width=True,
                key="radar_top",
            )
    except Exception as exc:  # noqa: BLE001
        st.info(f"雷达图渲染回退：{exc}")
        st.plotly_chart(
            build_radar_figure(list(DEFAULT_RADAR_SCORES)),
            use_container_width=True,
            key="radar_top_fb",
        )


def prompt_section_portrait(resume: str, jd: str) -> str:
    return f"""请用 **Markdown** 输出详尽「岗位深度画像」，聚焦：
- KPI 体系（指标名、口径、复盘节奏）
- **2026** 中国大陆薪酬区间（Base/奖金/总包假设写清城市与职级档位）
- 团队与协作动态（角色分工、跨部门、压力点）
要求：结构清晰、段落充分，避免泛泛而谈。

--- 简历 ---
{resume[:9500]}

--- JD ---
{jd[:9500]}
"""


def prompt_section_gap(resume: str, jd: str) -> str:
    return f"""请用 **Markdown** 输出详尽「核心差距与行动」：
- 对照 JD 与简历的 **详细清单**（技能/项目/行业/工具/证据）
- 每条差距给出 **可执行** 的补强动作与优先级
要求：清单尽量完整，可落地。

--- 简历 ---
{resume[:9500]}

--- JD ---
{jd[:9500]}
"""


def prompt_section_career(resume: str, jd: str) -> str:
    return f"""请用 **Markdown** 输出详尽「专属职业进路」：
- 2–4 条方向建议（结合简历 + 当前 JD）
- 每条说明适配理由、市场竞争力与风险
要求：具体、可决策。

--- 简历 ---
{resume[:9500]}

--- JD ---
{jd[:9500]}
"""


def prompt_interview_prediction(resume: str, jd: str) -> str:
    return f"""请基于以下「简历」与「岗位 JD」，输出 **恰好 10 道** Masterclass 级面试题与配套解析（「少而深」优于题量堆砌），整体达到「高管终面 / Bar Raiser」难度。

## 角色与标准
- 你代表 MBB 或一线科技大厂面试体系中的 **Senior HR Director**：每题追问要穿透；**「满分范例回答」是重头戏**，要比「硬核提问」篇幅更长、更细、更可背诵改写。
- 必须 **点名** 简历与 JD 中出现的公司、客户、品牌、平台、战役、职能（以材料中实际出现为准），禁止泛泛的「请介绍一个项目」。

## 每题固定结构（10 题全部遵守，用 Markdown）
使用二级标题：`## Qn · 简短题眼`（n=1..10），其下依次包含：

**[考察维度]**：只能是以下三者之一（写明中文 + 英文）：`领导力 Leadership` / `逻辑思维 Logical Thinking` / `行业认知 Industry Knowledge`（可附 1 句为何归为此类）。

**[硬核提问]**：1 段精炼但锋利的深度追问（仍须锚定材料中的具体项目/数据/组织情境）；可含「如果…你会如何…」类压力追问。**本段宜紧凑，为后面的范例回答留出篇幅。**

**[满分范例回答]**（**必须写深写透**，每题至少 **400–700 字** 量级，除非简历信息严重不足）：结构化、可背诵调整的高分回答。要求：
- 先给 **回答骨架**（3–6 个要点，编号或小标题）；
- 正文充分展开：融合 **量化指标 / 时间线 / 角色分工 / 风险与复盘 / 取舍理由**；自然嵌入 **1 种** 分析或表达框架（如 STAR、金字塔、Issue Tree、RACI、OKR、3C、利益相关方地图等，**择一即可，勿堆砌**）；
- 语气第一人称，专业、克制、可验证；若简历缺数据，明确写出「可补充的证据类型」与「如何在一周内补齐话术」。

## 版式与覆盖
- 题与题之间必须插入单独一行的 `---` 作为分割线。
- 10 题须在整体上覆盖：**项目深挖、方法论、客户与媒体关系、危机与舆情、数据与效果、跨团队冲突、战略取舍、职业规划与动机** 等维度（可一题多角，但避免问法重复）。
- **必须完整输出 Q1–Q10，缺一不可**；禁止输出到一半停止或省略某一题。

--- 简历 ---
{resume[:12000]}

--- JD ---
{jd[:12000]}
"""


def prompt_interview_debrief(transcript: str) -> str:
    return f"""你将对以下「面试复盘/转写」进行专业评估。请用 Markdown 输出，包含以下小节（使用二级标题）：

## 百分制评分
给出 0-100 的整数分数，并用 2-4 句说明扣分主要原因与加分点。

## 优点总结
条列 3-8 条，具体、可引用复盘中的表述。

## 详细改进建议
条列你认为需要改进的要点（可多于 8 条），每条包含：问题表现 → 改进方向 → 可练习的具体动作。

--- 面试复盘/转写 ---
{transcript[:14000]}
"""


def safe_section_markdown(
    client: OpenAI,
    label: str,
    user_prompt: str,
    *,
    max_tokens: int = 6144,
    system: str | None = None,
) -> str:
    try:
        text = chat_completion(
            client,
            system if system is not None else SYSTEM_HEADHUNTER,
            user_prompt,
            json_mode=False,
            max_tokens=max_tokens,
        )
        return text if text else f"（{label}：模型返回为空，请重试。）"
    except Exception as exc:  # noqa: BLE001
        return f"**{label} 生成失败**\n\n`{exc}`\n\n请检查网络或稍后重试。"


def run_full_deep_report(client: OpenAI, resume: str, jd: str) -> None:
    """雷达 + 三段正文，全部非流式；写入 session_state。"""
    scores, note = fetch_radar_scores_safe(client, resume, jd)
    st.session_state.radar_scores = scores
    st.session_state.radar_note = note

    st.session_state.deep_report_portrait = safe_section_markdown(
        client, "岗位深度画像", prompt_section_portrait(resume, jd), max_tokens=6144
    )
    st.session_state.deep_report_gap = safe_section_markdown(
        client, "核心差距与行动", prompt_section_gap(resume, jd), max_tokens=6144
    )
    st.session_state.deep_report_career = safe_section_markdown(
        client, "专属职业进路", prompt_section_career(resume, jd), max_tokens=6144
    )
    st.session_state.deep_report_ready = True


def ensure_session_state() -> None:
    if "resume_text" not in st.session_state:
        st.session_state.resume_text = ""
    if "jd_text" not in st.session_state:
        st.session_state.jd_text = ""
    if "radar_scores" not in st.session_state:
        st.session_state.radar_scores = None
    if "radar_note" not in st.session_state:
        st.session_state.radar_note = None
    if "deep_report_ready" not in st.session_state:
        st.session_state.deep_report_ready = False
    if "deep_report_portrait" not in st.session_state:
        st.session_state.deep_report_portrait = ""
    if "deep_report_gap" not in st.session_state:
        st.session_state.deep_report_gap = ""
    if "deep_report_career" not in st.session_state:
        st.session_state.deep_report_career = ""
    if "interview_prediction_md" not in st.session_state:
        st.session_state.interview_prediction_md = ""
    if "debrief_result_md" not in st.session_state:
        st.session_state.debrief_result_md = ""


def main() -> None:
    st.set_page_config(
        page_title="CareerVibe",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    ensure_session_state()
    client = get_client()

    gen_report = False

    with st.sidebar:
        st.markdown("### 材料")
        resume_file = st.file_uploader(
            "简历（PDF / Word）", type=["pdf", "docx"], key="resume_file"
        )
        jd_file = st.file_uploader("JD PDF（可选）", type=["pdf"], key="jd_pdf")
        jd_paste = st.text_area(
            "或粘贴 JD",
            height=120,
            placeholder="未上传 PDF 时在此粘贴…",
            key="jd_paste",
        )

        if resume_file is None:
            st.session_state.resume_text = ""
        else:
            try:
                st.session_state.resume_text = extract_resume_text(resume_file)
            except Exception as exc:  # noqa: BLE001 
                st.error(f"简历解析失败：{exc}")
                st.session_state.resume_text = ""

        if jd_file is not None:
            try:
                st.session_state.jd_text = extract_pdf_text(jd_file)
            except Exception as exc:  # noqa: BLE001
                st.error(f"JD 解析失败：{exc}")
                st.session_state.jd_text = ""
        elif jd_paste.strip():
            st.session_state.jd_text = jd_paste.strip()
        else:
            st.session_state.jd_text = ""

        st.divider()
        r_ok = len(st.session_state.resume_text) > 50
        j_ok = len(st.session_state.jd_text) > 50
        gen_report = st.button(
            "生成深度报告",
            type="primary",
            disabled=not (r_ok and j_ok),
            key="btn_deep_report",
        )

    st.title("CareerVibe")
    st.caption("深度岗位画像 · 五维雷达 · 三 Tab 报告 · 面试工坊（非流式）")

    materials_ok = len(st.session_state.resume_text) > 50 and len(
        st.session_state.jd_text
    ) > 50

    if materials_ok and gen_report:
        with st.spinner("正在生成深度报告（雷达 + 三段分析，请稍候）…"):
            run_full_deep_report(
                client,
                st.session_state.resume_text,
                st.session_state.jd_text,
            )

    if not materials_ok:
        st.info("请先在侧栏上传简历并填写 JD（文本需达到一定长度）。")
        return

    if not st.session_state.deep_report_ready:
        st.info("材料已就绪：点击侧栏「生成深度报告」，完成后将在此展示雷达与三个分析页签。")
        return

    scores = st.session_state.radar_scores
    if not isinstance(scores, list) or len(scores) != 5:
        scores = list(DEFAULT_RADAR_SCORES)
    render_radar_top(scores, overall_from_scores(scores))
    if st.session_state.radar_note:
        st.caption(st.session_state.radar_note)

    tab_a, tab_b, tab_c = st.tabs(
        ["岗位深度画像", "核心差距与行动", "专属职业进路"]
    )
    with tab_a:
        st.markdown(st.session_state.deep_report_portrait or "—")
    with tab_b:
        st.markdown(st.session_state.deep_report_gap or "—")
    with tab_c:
        st.markdown(st.session_state.deep_report_career or "—")

    st.divider()
    st.subheader("面试工坊")

    resume_body = st.session_state.resume_text
    jd_body = st.session_state.jd_text

    with st.expander("『面试预测』", expanded=False):
        st.caption(
            "MBB / 一线科技 Bar Raiser 风格：10 道 Masterclass + 考察维度 + 深度满分范例（非流式，一键生成）。"
        )
        run_pred = st.button(
            "运行预测（10 题）",
            key="btn_interview_predict_main",
            type="secondary",
        )
        if run_pred:
            with st.spinner("正在生成 10 题深度面试预测（每题范例较长，请耐心等待）…"):
                st.session_state.interview_prediction_md = safe_section_markdown(
                    client,
                    "面试预测",
                    prompt_interview_prediction(resume_body, jd_body),
                    max_tokens=8192,
                    system=SYSTEM_INTERVIEW_MASTER,
                )
        if st.session_state.interview_prediction_md:
            st.markdown(st.session_state.interview_prediction_md)

    with st.expander("『面试复盘』", expanded=False):
        st.caption("粘贴面试记录或复盘文字；点击运行后一次性输出评分与改进建议。")
        debrief_notes = st.text_area(
            "面试转写 / 复盘笔记",
            height=200,
            placeholder="题目、你的回答要点、面试官反馈、自我感觉等…",
            key="debrief_transcript_main",
        )
        run_debrief = st.button("运行复盘", key="btn_debrief_run_main")
        if run_debrief:
            if len(debrief_notes.strip()) < 30:
                st.warning("内容过短，请补充更多细节以便评分。")
            else:
                with st.spinner("正在评分与总结…"):
                    st.session_state.debrief_result_md = safe_section_markdown(
                        client,
                        "面试复盘",
                        prompt_interview_debrief(debrief_notes),
                        max_tokens=4096,
                    )
        if st.session_state.debrief_result_md:
            st.markdown(st.session_state.debrief_result_md)


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:  # noqa: BLE001
        try:
            st.error(f"应用异常：{exc}")
            st.code(traceback.format_exc())
        except Exception:
            raise exc from None
